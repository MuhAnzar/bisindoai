"""
Script untuk mengkonversi tabel pengujian blackbox dari Markdown ke Word (.docx)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def convert_markdown_to_word(md_file, output_file):
    """
    Konversi file markdown tabel pengujian ke Word
    """
    # Baca file markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Buat dokumen Word
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Judul
    title = doc.add_paragraph()
    title_run = title.add_run('Tabel Pengujian Black Box Testing')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Aplikasi BisindoCNNfi - Sistem Pembelajaran Bahasa Isyarat Indonesia')
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Parse tabel dari markdown
    lines = content.split('\n')
    in_table = False
    table_data = []
    
    for line in lines:
        if line.strip().startswith('| No |'):
            in_table = True
            # Header row
            headers = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(headers)
        elif in_table and line.strip().startswith('|---'):
            # Separator row - skip
            continue
        elif in_table and line.strip().startswith('|'):
            # Data row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Replace <br> dengan newline
            cells = [cell.replace('<br>', '\n') for cell in cells]
            table_data.append(cells)
        elif in_table and not line.strip().startswith('|'):
            # End of table
            break
    
    if table_data:
        # Buat tabel di Word
        num_rows = len(table_data)
        num_cols = len(table_data[0])
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                
                # Format header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Set background color for header
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '4472C4')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                else:
                    # Data rows
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                
                # Set cell borders
                set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, 
                               left={'val': 'single'}, right={'val': 'single'})
        
        # Adjust column widths
        widths = [Inches(0.4), Inches(1.2), Inches(1.5), Inches(1.5), Inches(1.8), Inches(2.0)]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    cell.width = widths[idx]
    
    # Add summary section
    doc.add_page_break()
    
    # Parse summary from markdown
    summary_started = False
    for line in lines:
        if '## Ringkasan Pengujian' in line:
            summary_started = True
            heading = doc.add_paragraph()
            heading_run = heading.add_run('Ringkasan Pengujian')
            heading_run.bold = True
            heading_run.font.size = Pt(14)
        elif summary_started:
            if line.strip().startswith('###'):
                # Sub-heading
                text = line.replace('###', '').strip()
                sub_heading = doc.add_paragraph()
                sub_heading_run = sub_heading.add_run(text)
                sub_heading_run.bold = True
                sub_heading_run.font.size = Pt(12)
            elif line.strip().startswith('-'):
                # Bullet point
                text = line.strip()[1:].strip()
                # Remove markdown bold
                text = text.replace('**', '')
                para = doc.add_paragraph(text, style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25)
            elif line.strip() and not line.strip().startswith('|'):
                # Normal text
                text = line.strip().replace('**', '')
                if text:
                    doc.add_paragraph(text)
    
    # Save document
    doc.save(output_file)
    print(f"✅ Dokumen Word berhasil dibuat: {output_file}")

if __name__ == "__main__":
    # Path files
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = r"C:\Users\Acer\.gemini\antigravity\brain\ad22f759-1233-4b96-9495-fb2603358300\tabel_pengujian_blackbox.md"
    output_file = os.path.join(script_dir, "UML", "Tabel_Pengujian_BlackBox.docx")
    
    # Pastikan folder UML ada
    os.makedirs(os.path.join(script_dir, "UML"), exist_ok=True)
    
    print("🔄 Mengkonversi Markdown ke Word...")
    convert_markdown_to_word(md_file, output_file)
    print(f"📄 File tersimpan di: {output_file}")
